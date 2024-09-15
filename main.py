from os import path, scandir
from docx import Document as docx_Document
from pymupdf import Document as pdf_Document
import json
from argparse import ArgumentParser, Namespace
from math import log10

from lexer import Lexer
from model import TermFreq


def index_document(doc_content: str) -> TermFreq:
    tf: TermFreq = dict()
    lexer = Lexer(doc_content)
    for term in lexer:
        if term in tf:
            tf[term] += 1
        else:
            tf[term] = 1
    return tf


def parse_pdf(doc_name: str) -> str:
    with open(doc_name, 'rb') as f:
        return "".join([page.get_text() for page in pdf_Document(f)])


def parse_docx(doc_name: str) -> str:
    with open(doc_name, 'rb') as f:
        return "\n".join([i.text for i in docx_Document(f).paragraphs])


def read_doc(doc_name: str) -> str | None:
    extension = path.splitext(doc_name)[-1]
    try:
        match extension:
            case ".docx":
                return parse_docx(doc_name)
            case '.pdf':
                return parse_pdf(doc_name)
            case _:
                return None
    except PermissionError as e:
        print(e)
        return None


type TermFreqIndex = dict[str, TermFreq]


def index(base_folder) -> None:
    tf_index: TermFreqIndex = dict()
    base_folder = path.abspath(base_folder)

    dir_list = [i.path for i in scandir(base_folder) if i.is_file()]

    for doc in dir_list:
        doc_content = read_doc(doc)
        if doc_content:
            print(f"Indexing \"{doc}\"")
            tf_index[doc] = index_document(doc_content)

    with open("index.json", 'w', encoding="utf-8") as w:
        json.dump(tf_index, w, ensure_ascii=False)


def tf(t: str, d: TermFreq) -> float:
    a = float(d[t] if t in d else 0)
    b = float(sum(d.values()))
    return a/b


def idf(t: str, d: TermFreqIndex) -> float:
    N: float = 1 + len(d)
    M: float = 1 + sum([1 if t in doc else 0 for doc in d.values()])
    return log10(N/M)


def search(prompt: str) -> None:
    with open("index.json", 'r', encoding="utf-8") as f:
        tf_index: TermFreqIndex = json.load(f)

    lexer = Lexer(prompt)
    result: dict[str, float] = dict()

    for doc in tf_index:
        # print(doc.split("\\")[-1])
        rank: float = 0
        for term in lexer:
            x = tf(term, tf_index[doc])*idf(term, tf_index)
            rank += x
            # print(term, x)
        result[doc.split("\\")[-1]] = rank

    result = dict(sorted(result.items(), key=lambda item: item[1],
                         reverse=True))

    for k, v in result.items():
        print(k, "\n\t", v)


def main() -> None:
    args_parser = ArgumentParser(prog="PROG")
    subparsers = args_parser.add_subparsers(title='Subcommands',
                                            required=True,
                                            dest="subparser_name")

    index_parser = subparsers.add_parser("index", help='Index the folder')
    index_parser.add_argument("folder", type=str, default="test",
                              help="The folder with documents to index")

    searh_index = subparsers.add_parser("search", help="Search the index")
    searh_index.add_argument("prompt", type=str, help="Search prompt")

    args: Namespace = args_parser.parse_args()

    match args.subparser_name:
        case "index":
            index(args.folder)
        case "search":
            search(args.prompt)
        case _:
            raise NotImplementedError


if __name__ == "__main__":
    main()
